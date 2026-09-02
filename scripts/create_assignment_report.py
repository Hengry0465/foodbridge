from __future__ import annotations

import json
import sys
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
EVIDENCE = ROOT / "docs" / "evidence"
OUTPUT = Path("/Users/easontan/Desktop/Assignment Requirement/FoodBridge Module 3 Assignment Report - Tan Tai Wei.docx")
TOC_JSON = ROOT / "docs" / "report-toc-pages.json"

SKILL_SCRIPTS = Path("/Users/easontan/.codex/plugins/cache/openai-primary-runtime/documents/26.826.12353/skills/documents/scripts")
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


INK = "18332D"
GREEN = "2F6B55"
GREEN_DARK = "214D3D"
GREEN_LIGHT = "E7F2EC"
BLUE = "315B73"
BLUE_LIGHT = "EAF1F5"
GOLD = "B7862F"
RED = "A23B3B"
MUTED = "586A64"
GRID = "A9BAB3"
LIGHT = "F5F8F6"
WHITE = "FFFFFF"
MONO = "Consolas"
BODY = "Arial"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GRID, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_paragraph_border(paragraph, color: str = GREEN, side: str = "left", size: str = "18", space: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), space)
    edge.set(qn("w:color"), color)
    borders.append(edge)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_field(run, instruction: str) -> None:
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_sep, display, fld_char_end])


def set_section_page_start(section, value: int) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(value))


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Title", 27, GREEN_DARK, 0, 12),
        ("Subtitle", 13, MUTED, 0, 18),
        ("Heading 1", 17, GREEN_DARK, 18, 8),
        ("Heading 2", 13, BLUE, 14, 6),
        ("Heading 3", 11, GREEN, 10, 4),
    ]:
        style = doc.styles[name]
        style.font.name = BODY
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    doc.styles["Heading 1"].paragraph_format.page_break_before = True

    caption = doc.styles["Caption"]
    caption.font.name = BODY
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.keep_with_next = False

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = BODY
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
        style.paragraph_format.space_after = Pt(3)

    if "Code" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = doc.styles["Code"]
    code_style.font.name = MONO
    code_style.font.size = Pt(7.4)
    code_style.font.color.rgb = rgb(INK)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), MONO)
    code_style.paragraph_format.left_indent = Inches(0.16)
    code_style.paragraph_format.right_indent = Inches(0.12)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(7)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.keep_together = True


def configure_document(doc: Document) -> None:
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "BMIT3173  |  FOODBRIDGE MODULE 3"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.runs[0]
    hr.font.name = BODY
    hr.font.size = Pt(7.5)
    hr.font.bold = True
    hr.font.color.rgb = rgb(GREEN)
    set_paragraph_border(hp, color=GREEN_LIGHT, side="bottom", size="8", space="3")

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Tan Tai Wei  |  2301297  •  ")
    r1.font.name = BODY
    r1.font.size = Pt(7.5)
    r1.font.color.rgb = rgb(MUTED)
    page_run = fp.add_run()
    set_field(page_run, "PAGE")
    page_run.font.name = BODY
    page_run.font.size = Pt(7.5)
    page_run.font.color.rgb = rgb(MUTED)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        set_paragraph_border(p, color=GREEN, side="bottom", size="12", space="5")


def add_paragraph(doc: Document, text: str = "", *, bold_prefix: str | None = None, italic: bool = False, keep: bool = False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        r = p.add_run(text)
        r.italic = italic
    p.paragraph_format.keep_together = keep
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.26)
    p.paragraph_format.first_line_indent = Inches(-0.16)


def add_callout(doc: Document, title: str, text: str, *, color: str = GREEN, fill: str = GREEN_LIGHT) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    title_run = p.add_run(title + "  ")
    title_run.bold = True
    title_run.font.color.rgb = rgb(color)
    p.add_run(text)
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, color=color, side="left", size="22", space="6")


def add_code(doc: Document, code: str, label: str | None = None) -> None:
    if label:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = rgb(BLUE)
    p = doc.add_paragraph(style="Code")
    p.add_run(code.strip("\n"))
    set_paragraph_shading(p, "F2F5F3")
    set_paragraph_border(p, color="B7C7C0", side="left", size="16", space="5")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], weights: list[float], *, font_size: float = 8.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, GREEN_DARK)
        set_cell_border(cell)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = BODY
                run.font.size = Pt(font_size)
                run.font.bold = True
                run.font.color.rgb = rgb(WHITE)
    set_repeat_table_header(table.rows[0])

    for row_idx, values in enumerate(rows, start=1):
        row = table.add_row()
        prevent_row_split(row)
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            if row_idx % 2 == 0:
                set_cell_shading(cell, LIGHT)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.02
                for run in p.runs:
                    run.font.name = BODY
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = rgb(INK)

    widths = column_widths_from_weights(weights, total_width_dxa=9780)
    apply_table_geometry(table, widths, table_width_dxa=9780, indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def crop_evidence() -> dict[str, Path]:
    crops = {
        "dashboard": ("02-recipient-dashboard.png", (0, 0, 720, 520)),
        "donations": ("03-available-donations.png", (0, 0, 720, 430)),
        "notification": ("04-match-notification.png", (0, 0, 720, 570)),
        "history": ("05-match-history.png", (0, 0, 720, 450)),
    }
    output: dict[str, Path] = {"landing": EVIDENCE / "01-foodbridge-landing.png"}
    for key, (name, crop) in crops.items():
        source = EVIDENCE / name
        target = EVIDENCE / f"report-{name}"
        with Image.open(source) as image:
            right = min(crop[2], image.width)
            bottom = min(crop[3], image.height)
            image.crop((crop[0], crop[1], right, bottom)).save(target)
        output[key] = target
    output["entity"] = EVIDENCE / "entity-class-diagram.png"
    output["observer"] = EVIDENCE / "observer-pattern-diagram.png"
    return output


def add_figure(doc: Document, path: Path, caption: str, *, width: float = 6.45, max_height: float = 6.7) -> None:
    with Image.open(path) as img:
        ratio = img.height / img.width
    final_width = width
    final_height = final_width * ratio
    if final_height > max_height:
        final_height = max_height
        final_width = max_height / ratio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(path), width=Inches(final_width), height=Inches(final_height))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption.split(".", 1)[0])
    cap = doc.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_with_next = False


def add_toc(doc: Document) -> None:
    pages = {}
    if TOC_JSON.exists():
        pages = json.loads(TOC_JSON.read_text())
    entries = [
        ("Declaration and AI-use disclosure", "declaration", 0),
        ("Executive summary", "executive", 0),
        ("1. Introduction", "1. Introduction", 0),
        ("2. Module Description", "2. Module Description", 0),
        ("3. Entity Classes", "3. Entity Classes", 0),
        ("4. Design Pattern — Observer", "4. Design Pattern", 0),
        ("5. Software Security", "5. Software Security", 0),
        ("6. Web Services and IFA", "6. Web Services", 0),
        ("7. References", "7. References", 0),
        ("Appendix A — Requirement Compliance Matrix", "Appendix A", 0),
        ("Appendix B — Verification Evidence", "Appendix B", 0),
    ]
    title = doc.add_paragraph("Table of Contents")
    title.style = doc.styles["Heading 1"]
    title.paragraph_format.page_break_before = False
    set_paragraph_border(title, color=GREEN, side="bottom", size="12", space="5")
    add_paragraph(doc, "Page numbers refer to the final rendered document. Heading styles are retained so Word can also generate a native TOC if required.", italic=True)
    for label, key, level in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22 * level)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.4), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(label)
        r.font.name = BODY
        r.font.size = Pt(10.5)
        if level == 0:
            r.bold = True
            r.font.color.rgb = rgb(INK)
        p.add_run("\t" + str(pages.get(key, "—")))


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge = p.add_run("BMIT3173")
    badge.font.name = BODY
    badge.font.size = Pt(12)
    badge.font.bold = True
    badge.font.color.rgb = rgb(WHITE)
    set_paragraph_shading(p, GREEN_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(38)
    r = p.add_run("INTEGRATIVE PROGRAMMING")
    r.font.name = BODY
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = rgb(GREEN)
    r.font.letter_spacing = Pt(1.2)

    p = doc.add_paragraph("Assignment Report", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p = doc.add_paragraph("FoodBridge — Module 3: Request & Auto Matching", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_callout(
        doc,
        "SDG 2 · ZERO HUNGER",
        "A secure food-request and automatic-matching module that connects available donations with recipients while prioritising expiry proximity, fairness, and transparent notifications.",
        color=GREEN_DARK,
        fill=GREEN_LIGHT,
    )

    table = doc.add_table(rows=0, cols=2)
    metadata = [
        ("Student Name", "Tan Tai Wei"),
        ("Student ID", "2301297"),
        ("Programme", "Not provided in the supplied materials"),
        ("Tutorial Group", "RIS3S1G2 / RSD2S3G3"),
        ("Group", "Group 4"),
        ("Assignment Session", "202605"),
        ("Module", "Module 3 — Request & Auto Matching"),
        ("Design Pattern", "Observer"),
        ("Submission Date", "28 August 2026"),
    ]
    for idx, (key, value) in enumerate(metadata):
        row = table.add_row()
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_border(cell, color="D4DFDA", size="4")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[0].text = key
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "EEF4F1")
        if idx % 2 == 1:
            set_cell_shading(row.cells[1], "FAFBFA")
        for run in row.cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = rgb(GREEN_DARK)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = BODY
                    run.font.size = Pt(9.4)
    set_repeat_table_header(table.rows[0])
    apply_table_geometry(table, [2500, 7280], table_width_dxa=9780, indent_dxa=120)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared from the verified Laravel 13.20.0 implementation")
    r.font.name = BODY
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = rgb(MUTED)
    p.add_run().add_break(WD_BREAK.PAGE)


def build_report() -> None:
    assets = crop_evidence()
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "FoodBridge Module 3 Assignment Report"
    doc.core_properties.subject = "BMIT3173 Integrative Programming — Request & Auto Matching"
    doc.core_properties.author = "Tan Tai Wei (2301297)"
    doc.core_properties.keywords = "FoodBridge, Observer, Laravel, SDG 2, security, IFA"
    doc.core_properties.comments = "Generated from the verified FoodBridge Module 3 implementation and assignment requirements."

    add_cover(doc)

    add_heading(doc, "Declaration and AI-use disclosure", 1)
    doc.paragraphs[-1].paragraph_format.page_break_before = False
    add_paragraph(doc, "I declare that this report describes the FoodBridge Module 3 implementation submitted under my student identity. The implementation evidence, class paths, screenshots, test results, and interface definitions in this report were verified against the local project before submission.")
    add_paragraph(doc, "Generative AI was used as an implementation and writing assistant to review the supplied assignment brief, improve code quality, generate diagrams, organise evidence, and draft this report. The student remains responsible for reviewing the final work, understanding the implementation, and ensuring compliance with institutional academic-integrity rules.")
    add_callout(doc, "Student review required", "Before submission, confirm the programme name on the cover and add a signature/date if your lecturer requires a signed declaration.", color=GOLD, fill="FFF7E7")

    add_heading(doc, "Executive summary", 1)
    doc.paragraphs[-1].paragraph_format.page_break_before = False
    add_paragraph(doc, "FoodBridge supports Sustainable Development Goal 2 (Zero Hunger) by coordinating surplus-food donations with eligible recipients. This report covers Module 3, Request & Auto Matching. A recipient submits a request for a category, quantity, and pickup time; the module verifies identity, checks eligible donation availability, allocates stock fairly, records each allocation, and informs affected users.")
    add_paragraph(doc, "The completed implementation uses the Observer design pattern. AutoMatchingService acts as the subject client, MatchPublisher maintains observers, and DonorNotifier and RecipientNotifier react to MatchSucceeded, PartialMatch, or MatchFailed outcomes. The pattern is distinct from MVC and keeps the matching algorithm independent from notification behaviour.")
    add_paragraph(doc, "Security is treated as an implementation concern rather than a descriptive add-on. Database transactions and row locks prevent concurrent over-allocation. Inter-module APIs require HMAC-SHA256 signatures, fresh timestamps, unique request IDs, payload fingerprints, consistent response envelopes, and a rate limit of 10 requests per minute per module/IP. Mandatory input validation and authenticated ownership checks provide additional protection. Automated verification completed successfully: 14 tests and 52 assertions passed, Laravel Pint passed, Composer validation passed, and the end-to-end browser flow produced a successful 12-of-12 match.")

    doc.add_page_break()
    add_toc(doc)

    add_heading(doc, "1. Introduction", 1)
    add_heading(doc, "1.1 System overview", 2)
    add_paragraph(doc, "FoodBridge is a modular web application for coordinating surplus food between donors and recipients. Module 3 owns food requests, automatic matching records, and match notifications. It consumes identity and donation-availability information from other modules and exposes request/match information to collaborating modules through a REST/JSON interface.")
    add_figure(doc, assets["landing"], "Figure 1. FoodBridge public landing page. View: resources/views/welcome.blade.php; route: routes/web.php.", width=6.35, max_height=4.5)

    add_heading(doc, "1.2 Contribution to SDG 2: Zero Hunger", 2)
    add_paragraph(doc, "United Nations Sustainable Development Goal 2 seeks to end hunger and improve access to sufficient, nutritious food, particularly for vulnerable people (United Nations, n.d.). FoodBridge contributes operationally by reducing the time between donation availability and recipient allocation, favouring food that expires sooner, supporting partial fulfilment when stock is limited, and retaining a transparent request history.")
    add_bullet(doc, "Donors obtain a structured channel for offering usable surplus food.")
    add_bullet(doc, "Recipients can request food by category and quantity without manually coordinating with multiple donors.")
    add_bullet(doc, "Expiry-aware matching helps reduce avoidable food waste.")
    add_bullet(doc, "Notification records make allocation outcomes visible to both sides.")

    doc.add_page_break()
    add_heading(doc, "1.3 Target users and module boundary", 2)
    add_table(doc, ["Actor / Module", "Need", "Module 3 responsibility"], [
        ["Recipient", "Request suitable food and know the result", "Create general or donation-specific requests; view status/history; withdraw only unmatched pending requests"],
        ["Donor", "Know when donated stock is allocated", "Receive a donor notification for every match involving the donor's donation"],
        ["Module 1 — Identity", "Authoritative user identity and role", "Consumed by Module 3 before matching; local recipient verification is a development fallback"],
        ["Module 2 — Donation", "Authoritative eligible stock", "Consumed by Module 3; results constrain local candidates by category and availability"],
        ["Other modules", "Read request/match outcomes", "Use authenticated REST endpoints and the common IFA response envelope"],
    ], [1.4, 2.2, 3.5], font_size=8.5)
    add_callout(doc, "Module boundary", "Module 3 owns requests and matches. It does not create donor identity records or originate donation listings; those are integration responsibilities of Modules 1 and 2.")

    add_heading(doc, "2. Module Description", 1)
    add_heading(doc, "2.1 Complete functions and class paths", 2)
    add_table(doc, ["Function", "User outcome", "Primary implementation path", "Verification"], [
        ["Authenticated recipient access", "Only a recipient session can create a browser request", "app/Http/Middleware/EnsureRole.php; routes/web.php", "Feature tests and live login"],
        ["Browse donation availability", "View category, remaining quantity, expiry, donor, and pickup point", "app/Http/Controllers/FoodRequestController.php; resources/views/dashboard.blade.php", "UI Figure 2"],
        ["Submit general request", "Request a category, quantity, and future pickup time", "app/Http/Requests/StoreFoodRequest.php; FoodRequestController::store()", "Validation and matching tests"],
        ["Request selected donation", "Prefer a chosen eligible donation while allowing remaining quantity to be filled from others", "FoodRequest::preferredDonation(); AutoMatchingService::match()", "AutoMatchingTest"],
        ["Automatic match", "Allocate by category, availability, expiry proximity, creation order, and ID; allow partial matches", "app/Services/AutoMatchingService.php", "14-test suite"],
        ["Observer notification", "Recipient and each affected donor receive an outcome message", "app/Services/MatchPublisher.php; app/Observers/*Notifier.php", "MatchPublisherTest and feature tests"],
        ["Request history", "See requested, matched, pending, partial, or withdrawn states", "FoodRequestController::dashboard(); resources/views/dashboard.blade.php", "UI Figure 4"],
        ["Withdraw unmatched request", "Withdraw only a request still in pending state and owned by the user", "FoodRequestController::destroy(); Api/FoodRequestController::destroy()", "Ownership/status tests"],
    ], [1.4, 1.7, 2.3, 1.25], font_size=7.55)

    add_heading(doc, "2.2 Recipient dashboard and request creation", 2)
    add_paragraph(doc, "The dashboard derives the recipient from the authenticated session. The browser form does not accept a trusted hidden recipient identifier; the controller assigns recipient_id from the authenticated user. A recipient may submit a general request or select an available donation. Category consistency, unexpired stock, positive availability, quantity range, and future pickup time are validated before a request is created.")
    add_figure(doc, assets["dashboard"], "Figure 2. Authenticated recipient dashboard and general request form. Controller: app/Http/Controllers/FoodRequestController.php; view: resources/views/dashboard.blade.php.", width=6.2, max_height=4.7)
    add_figure(doc, assets["donations"], "Figure 3. Available donations and donation-specific request action. Model: app/Models/Donation.php; matching service: app/Services/AutoMatchingService.php.", width=6.2, max_height=4.5)

    add_heading(doc, "2.3 Automatic matching workflow", 2)
    add_paragraph(doc, "After creation, the controller invokes AutoMatchingService inside a database transaction. The service first verifies the recipient through Module 1 (or the local development fallback) and requests eligible donation IDs from Module 2 when its URL is configured. It then queries local donations using the following deterministic order:")
    for item in [
        "same category as the request;",
        "status is available, quantity_available is greater than zero, and expires_at is in the future;",
        "preferred donation first when one is selected;",
        "earliest expiry first to reduce waste;",
        "earliest creation time, then lowest ID, as stable FCFS tie-breakers; and",
        "partial allocation across multiple donations until the requested quantity is satisfied or eligible stock is exhausted.",
    ]:
        add_bullet(doc, item)
    add_code(doc, """$donations = Donation::query()
    ->where('category', $request->category)
    ->where('status', 'available')
    ->where('quantity_available', '>', 0)
    ->where('expires_at', '>', now())
    ->when($request->preferred_donation_id,
        fn ($query) => $query->orderByRaw(
            'CASE WHEN id = ? THEN 0 ELSE 1 END',
            [$request->preferred_donation_id]
        ))
    ->orderBy('expires_at')->orderBy('created_at')->orderBy('id')
    ->lockForUpdate()->get();""", "Code excerpt 1. Eligibility, preference, expiry proximity, FCFS tie-breakers, and row locking — app/Services/AutoMatchingService.php")

    add_heading(doc, "2.4 Match outcome and notification", 2)
    add_paragraph(doc, "The saved request status is matched when the full quantity is allocated, partial when at least one portion but not the full quantity is allocated, and pending when no eligible stock is found. Exactly one concrete outcome event is published. Observers persist separate notification records for the recipient and each unique affected donor.")
    add_figure(doc, assets["notification"], "Figure 4. Successful 12-of-12 allocation and recipient Observer notification. Observer: app/Observers/RecipientNotifier.php.", width=5.25, max_height=4.2)
    add_figure(doc, assets["history"], "Figure 5. Request history after matching, with the allocated quantity and status. View: resources/views/dashboard.blade.php.", width=5.25, max_height=3.4)

    doc.add_page_break()
    add_heading(doc, "2.5 Separation of concerns", 2)
    add_table(doc, ["Layer", "Responsibility", "Representative classes"], [
        ["Presentation", "Forms, donor cards, notifications, and history", "resources/views/welcome.blade.php; auth/login.blade.php; dashboard.blade.php"],
        ["HTTP / application", "Route protection, validation, orchestration, IFA responses", "routes/web.php; routes/api.php; controllers; StoreFoodRequest"],
        ["Domain services", "Matching, inter-module signing/client, Observer publication", "AutoMatchingService; ModuleApiClient; ModuleRequestSigner; MatchPublisher"],
        ["Domain model", "Object references, casts, persistence intent", "User; Donation; FoodRequest; MatchRecord; MatchNotification"],
        ["Infrastructure", "Database transactions, middleware, rate limiter, configuration", "VerifyModuleRequest; AppServiceProvider; migrations; config/integrations.php"],
    ], [1.2, 2.5, 3.0], font_size=8.2)

    add_heading(doc, "3. Entity Classes", 1)
    add_heading(doc, "3.1 Entity class diagram", 2)
    add_paragraph(doc, "Figure 6 is an entity class diagram rather than an ERD. It expresses navigable object references and multiplicities used in the application. Database columns still implement referential integrity, but business code navigates Eloquent objects such as $request->recipient, $match->donation, and $notification->foodRequest instead of manually passing foreign-key values between services.")
    add_figure(doc, assets["entity"], "Figure 6. Entity class diagram for FoodBridge Module 3. Source classes: app/Models/*.php.", width=6.35, max_height=4.4)

    add_heading(doc, "3.2 Entity responsibilities and object references", 2)
    add_table(doc, ["Entity class", "Core state", "Object references", "Purpose"], [
        ["User", "name, email, role", "donations: Donation[*]; foodRequests: FoodRequest[*]; matchNotifications: MatchNotification[*]", "Represents recipient or donor identity and navigates owned records"],
        ["Donation", "foodName, category, quantityAvailable, expiresAt, status", "donor: User; matches: MatchRecord[*]", "Represents eligible stock supplied by a donor"],
        ["FoodRequest", "category, quantityRequested, quantityMatched, pickup time, status", "recipient: User; preferredDonation: Donation?; matches: MatchRecord[*]; notifications: MatchNotification[*]", "Aggregate root for a recipient's requested quantity and outcome"],
        ["MatchRecord", "quantityAllocated, status", "foodRequest: FoodRequest; donation: Donation", "Allocation line connecting one request to one donation"],
        ["MatchNotification", "type, message, readAt", "user: User; foodRequest: FoodRequest", "Persisted Observer reaction addressed to a recipient or donor"],
    ], [1.15, 1.75, 2.65, 2.0], font_size=7.8)

    add_heading(doc, "3.3 Consistent implementation example", 2)
    add_code(doc, """class FoodRequest extends Model
{
    public function recipient(): BelongsTo
    {
        return $this->belongsTo(User::class, 'recipient_id');
    }

    public function preferredDonation(): BelongsTo
    {
        return $this->belongsTo(Donation::class, 'preferred_donation_id');
    }

    public function matches(): HasMany
    {
        return $this->hasMany(MatchRecord::class, 'request_id');
    }

    public function notifications(): HasMany
    {
        return $this->hasMany(MatchNotification::class, 'request_id');
    }
}""", "Code excerpt 2. Object-reference relationships — app/Models/FoodRequest.php")
    add_paragraph(doc, "These relationships match the diagram and are reused by the matching service, controllers, views, and observers. The implementation therefore avoids a diagram-only domain model.")

    add_heading(doc, "4. Design Pattern — Observer", 1)
    add_heading(doc, "4.1 Pattern description", 2)
    add_paragraph(doc, "The Observer pattern defines a one-to-many dependency so that when a subject changes state, registered observers are notified (Gamma et al., 1994). In Module 3, the important change is a completed matching attempt. The matching algorithm publishes one immutable outcome object, while notification components respond independently. This is a behavioural pattern and is separate from the Laravel MVC request-processing structure.")
    add_figure(doc, assets["observer"], "Figure 7. Observer implementation class diagram. Source: app/Contracts, app/Events, app/Services, and app/Observers.", width=6.4, max_height=4.05)

    add_heading(doc, "4.2 Participant mapping", 2)
    add_table(doc, ["Observer role", "Module 3 class", "Responsibility"], [
        ["Observer interface", "App\\Contracts\\MatchObserver", "Declares update(MatchOutcome): void"],
        ["Subject / publisher", "App\\Services\\MatchPublisher", "Maintains an observer collection; attach, detach, and notify"],
        ["Subject client", "App\\Services\\AutoMatchingService", "Selects MatchSucceeded, PartialMatch, or MatchFailed after commit"],
        ["Concrete observer", "App\\Observers\\RecipientNotifier", "Creates one request-outcome notification for the recipient"],
        ["Concrete observer", "App\\Observers\\DonorNotifier", "Creates a match notification for each unique affected donor"],
        ["Event hierarchy", "App\\Events\\MatchOutcome and subclasses", "Carries the matched FoodRequest and supplies the outcome type"],
        ["Composition root", "App\\Providers\\AppServiceProvider", "Creates the singleton publisher and attaches both observers"],
    ], [1.35, 2.45, 3.3], font_size=8.0)

    add_heading(doc, "4.3 Implementation", 2)
    add_code(doc, """interface MatchObserver
{
    public function update(MatchOutcome $event): void;
}

class MatchPublisher
{
    private array $observers = [];

    public function attach(MatchObserver $observer): void
    {
        $this->observers[$observer::class] = $observer;
    }

    public function notify(MatchOutcome $event): void
    {
        foreach ($this->observers as $observer) {
            $observer->update($event);
        }
    }
}""", "Code excerpt 3. Observer interface and publisher — app/Contracts/MatchObserver.php; app/Services/MatchPublisher.php")
    add_code(doc, """$event = match ($outcome) {
    'matched' => new MatchSucceeded($request),
    'partial' => new PartialMatch($request),
    default => new MatchFailed($request),
};
$this->matchPublisher->notify($event);""", "Code excerpt 4. Outcome publication — app/Services/AutoMatchingService.php")

    add_heading(doc, "4.4 Justification", 2)
    add_paragraph(doc, "Observer is suitable because matching and notification change for different reasons. The matching service should decide allocations; it should not know how many channels, recipient messages, donor messages, audit handlers, or analytics handlers exist. The publisher depends only on MatchObserver, which provides the following benefits:")
    add_bullet(doc, "Open/Closed Principle: an EmailNotifier or AuditObserver can be added by attachment without editing AutoMatchingService.")
    add_bullet(doc, "Single Responsibility Principle: matching remains focused on identity checks, candidate ordering, and allocation; observers focus on reactions.")
    add_bullet(doc, "Testability: MatchPublisherTest verifies attach, notify, and detach independently from the database algorithm.")
    add_bullet(doc, "Explicit outcomes: matched, partial, and pending states are represented by different event types instead of scattered conditionals.")
    add_callout(doc, "Why not Singleton or MVC?", "The publisher is registered once by the service container, but the assessed pattern is its one-to-many notification collaboration. MVC is the framework structure and is not claimed as the unique design pattern.", color=BLUE, fill=BLUE_LIGHT)

    add_heading(doc, "5. Software Security", 1)
    add_heading(doc, "5.1 Security posture", 2)
    add_paragraph(doc, "The controls address integrity, authentication, authorisation, availability, and safe failure. The current OWASP Top 10 identifies broken access control, cryptographic failures, insecure design, and mishandling of exceptional conditions as important web risks (OWASP Foundation, 2025). The two assessed secure-coding strategies below are distinct and do not count input validation as either strategy, in accordance with the assignment brief.")
    add_table(doc, ["Threat", "Attack / failure scenario", "Impact", "Primary secure-coding strategy"], [
        ["Concurrent double allocation", "Two requests read the same quantity before either update commits", "Negative or overstated stock; conflicting pickup promises; loss of integrity", "Transaction + pessimistic row lock + atomic allocation and constraints"],
        ["Forged or replayed module API request", "An attacker changes the payload, reuses a captured request, or impersonates another module", "Unauthorised request creation/withdrawal; stock depletion; unreliable audit trail", "HMAC-signed canonical request + fresh timestamp + request ID/fingerprint idempotency"],
    ], [1.35, 2.4, 1.9, 2.5], font_size=8.1)

    add_heading(doc, "5.2 Threat 1 — race condition and double allocation", 2)
    add_paragraph(doc, "A naïve read-modify-write flow is unsafe under concurrency. Suppose donation D has 10 portions and two recipients concurrently request 8. If both processes read 10 before either writes, each may allocate 8, promising 16 portions from stock that only contains 10. Even correct form validation cannot prevent this timing flaw. The result violates integrity and may leave donors and recipients with inconsistent commitments.")
    add_heading(doc, "5.2.1 Strategy — transactional pessimistic locking", 3)
    add_paragraph(doc, "The service executes allocation in DB::transaction() and retrieves candidate rows with lockForUpdate(). The database holds selected donation rows exclusively until the transaction commits or rolls back. Each allocation uses the current locked quantity, creates a MatchRecord, decrements quantity_available, increments version, and changes a depleted donation to reserved. The outer controller transaction also makes request creation and matching one unit of work. Laravel recommends placing pessimistic locks inside a transaction so selected data remains unmodified during the operation (Laravel, 2026a).")
    add_code(doc, """$outcome = DB::transaction(function () use ($request) {
    $donations = Donation::query()
        ->where('quantity_available', '>', 0)
        ->lockForUpdate()->get();

    foreach ($donations as $donation) {
        $allocated = min($remaining, $donation->quantity_available);
        MatchRecord::create([
            'request_id' => $request->id,
            'donation_id' => $donation->id,
            'quantity_allocated' => $allocated,
        ]);
        $donation->quantity_available -= $allocated;
        $donation->version++;
        $donation->save();
    }
}, 3);""", "Code excerpt 5. Atomic allocation under a row lock — app/Services/AutoMatchingService.php")
    add_callout(doc, "Security result", "A competing transaction must wait and then observe the already-reduced quantity. If any operation fails, the allocation records and stock changes roll back together.")

    add_heading(doc, "5.3 Threat 2 — forged and replayed inter-module requests", 2)
    add_paragraph(doc, "An internal API is not automatically trusted. Without message authentication, a caller could forge a recipient request, change a quantity in transit, or impersonate Module 2. A valid request could also be captured and replayed to create multiple identical rows. Request IDs alone are insufficient if the attacker changes the payload associated with an ID.")
    add_heading(doc, "5.3.1 Strategy — HMAC authentication, freshness, and idempotency", 3)
    add_paragraph(doc, "Every API request includes X-Module-ID, X-Request-ID, X-Timestamp, and X-Signature. ModuleRequestSigner builds a canonical string from the module ID, request ID, timestamp, uppercase method, exact request target, and SHA-256 body hash, then calculates HMAC-SHA256 using the shared module secret. VerifyModuleRequest rejects missing/invalid fields, timestamps outside the five-minute clock-skew window, and signatures that fail constant-time hash_equals comparison.")
    add_code(doc, """$canonical = implode("\\n", [
    $moduleId, $requestId, $timestamp,
    strtoupper($method), $requestTarget,
    hash('sha256', $body),
]);

return hash_hmac('sha256', $canonical, $secret);""", "Code excerpt 6. Canonical HMAC-SHA256 signature — app/Services/ModuleRequestSigner.php")
    add_paragraph(doc, "For POST /requests, the API stores client_request_id and a SHA-256 fingerprint of the security-relevant body. Repeating the same ID and same fingerprint returns status duplicate without creating another row. Reusing that ID with a different payload returns HTTP 409 idempotency_conflict. The request_fingerprint attribute is hidden from JSON output. Rate limiting allows 10 requests per minute per module/IP, reducing brute-force and denial-of-service pressure.")
    add_code(doc, """if ($existing = FoodRequest::where('client_request_id', $clientId)->first()) {
    if (! hash_equals($existing->request_fingerprint, $fingerprint)) {
        return $this->error(
            'The requestID has already been used for a different payload.',
            409,
            'idempotency_conflict'
        );
    }
    return $this->respond($existing, 200, 'duplicate');
}""", "Code excerpt 7. Replay-safe idempotency — app/Http/Controllers/Api/FoodRequestController.php")

    add_heading(doc, "5.4 Mandatory input validation and supporting controls", 2)
    add_paragraph(doc, "Input validation is implemented as a mandatory baseline but is not counted as either secure-coding strategy above. StoreFoodRequest applies allow-listed categories, integer bounds, future pickup time, role-aware recipient existence, and donation eligibility. Browser requests prohibit client-supplied recipient_id; API requests require it and verify that it belongs to a recipient. Additional controls include session authentication, role middleware, ownership checks for browser withdrawal, ORM queries, password hashing, CSRF protection on web forms, hidden sensitive attributes, uniform API errors, and environment-based secrets.")
    add_code(doc, """'category' => ['required', 'string',
    'in:Cooked Meals,Bakery,Fresh Produce,Packaged Goods'],
'quantity' => ['required', 'integer', 'min:1', 'max:10000'],
'preferred_pickup_at' => ['required', 'date', 'after:now'],
'recipient_id' => [
    $apiRequest ? 'required' : 'prohibited',
    'integer',
    Rule::exists('users', 'id')->where(
        fn ($query) => $query->where('role', 'recipient')
    ),
],""", "Code excerpt 8. Mandatory request validation — app/Http/Requests/StoreFoodRequest.php")

    add_heading(doc, "6. Web Services and Interoperability Framework Agreement", 1)
    add_heading(doc, "6.1 Overview", 2)
    add_paragraph(doc, "Module 3 both exposes and consumes REST/JSON web services. All module-to-module traffic uses the same interoperability framework agreement (IFA): request identification, ISO-8601 timestamps, HMAC authentication, JSON, explicit HTTP status codes, and a response envelope containing requestID, timestamp, and status. Service secrets and base URLs are read from environment configuration.")

    add_heading(doc, "6.2 Exposed service mechanism", 2)
    add_table(doc, ["Method", "URL", "Function", "Success", "Typical consumers"], [
        ["POST", "/api/v1/requests", "Create a request and immediately run automatic matching", "201 success; 200 duplicate", "Module 2, integration client"],
        ["GET", "/api/v1/requests/{foodRequest}/match", "Read request status and allocation records", "200 success", "Pickup/admin/reporting modules"],
        ["DELETE", "/api/v1/requests/{foodRequest}", "Withdraw only an unmatched pending request", "200 success; 409 otherwise", "Recipient-facing or coordination module"],
    ], [0.9, 1.85, 2.45, 1.25, 1.55], font_size=7.8)
    add_paragraph(doc, "Route implementation: routes/api.php and app/Http/Controllers/Api/FoodRequestController.php. VerifyModuleRequest is prepended to the API middleware group so authentication occurs before route-model binding. The named module-api throttle is applied to the v1 group.")

    add_heading(doc, "6.3 Request IFA", 2)
    add_heading(doc, "6.3.1 Mandatory request headers", 3)
    add_table(doc, ["Field", "Type / format", "Mandatory", "Validation", "Purpose"], [
        ["X-Module-ID", "String, 2–40: A–Z a–z 0–9 . _ -", "Yes", "Allow-list regex", "Identifies the calling module and rate-limit bucket"],
        ["X-Request-ID", "String, 8–100: A–Z a–z 0–9 . _ : -", "Yes", "Allow-list regex; unique for mutation", "Traceability and idempotency"],
        ["X-Timestamp", "ISO-8601 datetime", "Yes", "Parsable and within ±300 seconds", "Freshness and replay resistance"],
        ["X-Signature", "64-character hexadecimal HMAC-SHA256", "Yes", "Recalculated and compared with hash_equals", "Authenticity and integrity"],
        ["Accept", "application/json", "Yes for client", "HTTP content negotiation", "Response representation"],
        ["Content-Type", "application/json", "POST", "Request parsing", "Body representation"],
    ], [1.2, 2.0, 1.05, 2.2, 2.1], font_size=7.45)
    add_heading(doc, "6.3.2 POST request body", 3)
    add_table(doc, ["Field", "Type", "Mandatory", "Rule", "Meaning"], [
        ["recipient_id", "Integer", "Yes", "Existing User with role recipient", "Recipient requesting food"],
        ["donation_id", "Integer / null", "No", "Eligible, available, unexpired and category-compatible", "Preferred donation; not exclusive"],
        ["category", "String enum", "Yes", "Cooked Meals | Bakery | Fresh Produce | Packaged Goods", "Requested food category"],
        ["quantity", "Integer", "Yes", "1–10,000", "Requested portions"],
        ["preferred_pickup_at", "ISO-8601 datetime", "Yes", "Future datetime", "Requested pickup time"],
    ], [1.3, 1.1, 0.85, 2.8, 1.7], font_size=7.8)
    add_code(doc, """POST /api/v1/requests HTTP/1.1
X-Module-ID: module-2
X-Request-ID: module2-demo-20260828-002
X-Timestamp: 2026-08-28T22:22:50+08:00
X-Signature: <HMAC-SHA256 hex>
Content-Type: application/json

{
  "recipient_id": 2,
  "category": "Packaged Goods",
  "quantity": 12,
  "preferred_pickup_at": "2026-08-29T10:00:00+08:00"
}""", "IFA example 1. Signed request")

    add_heading(doc, "6.4 Response IFA", 2)
    add_table(doc, ["Field", "Type", "Mandatory", "Values / rule", "Purpose"], [
        ["requestID", "String / null", "Yes", "Echoes validated request ID; null only when absent/invalid", "Correlates response with request"],
        ["timestamp", "ISO-8601 datetime", "Yes", "Server response time", "Ordering, tracing, and freshness"],
        ["status", "String", "Yes", "success | duplicate | idempotency_conflict | error", "Machine-readable outcome"],
        ["data", "Object", "Success", "FoodRequest and optional matches/donations", "Business result"],
        ["message", "String", "Error", "Safe client-facing explanation", "Failure detail without stack trace"],
        ["errors", "Object", "Validation error", "Field-keyed messages", "Actionable 422 response"],
    ], [1.15, 1.4, 1.0, 2.25, 2.4], font_size=7.6)
    add_code(doc, """HTTP/1.1 201 Created
Content-Type: application/json

{
  "requestID": "module2-demo-20260828-002",
  "timestamp": "2026-08-28T22:22:50+08:00",
  "status": "success",
  "data": {
    "id": "<request-uuid>",
    "recipient_id": 2,
    "category": "Packaged Goods",
    "quantity_requested": 12,
    "quantity_matched": 12,
    "status": "matched"
  }
}""", "IFA example 2. Verified success envelope; internal request fingerprint is hidden")
    add_table(doc, ["HTTP code", "Status", "Meaning"], [
        ["200", "success / duplicate", "Read, withdrawal, or idempotent repeat completed"],
        ["201", "success", "Request created and matching completed"],
        ["400", "error", "Malformed request identifier or timestamp"],
        ["401", "error", "Module ID/signature missing or invalid"],
        ["408", "error", "Timestamp outside allowed clock-skew window"],
        ["409", "error / idempotency_conflict", "Invalid state transition or reused ID with different payload"],
        ["422", "error", "Body validation or recipient verification failed"],
        ["429", "error", "Rate limit exceeded"],
        ["500 / 503", "error", "Safe generic service/configuration failure"],
    ], [1.1, 2.2, 4.6], font_size=8.0)

    add_heading(doc, "6.5 Consumed services", 2)
    add_table(doc, ["Provider", "Method / URL", "Request", "Expected response", "Timeout / retry", "Fallback"], [
        ["Module 1 — Identity", "GET {IDENTITY_URL}/api/v1/users/{id}", "Signed headers; recipient ID in path", "status, timestamp, data.id, data.role=recipient", "3 seconds; one retry (two attempts total)", "Local User role check when URL is not configured"],
        ["Module 2 — Donation", "GET {DONATION_URL}/api/v1/donations", "Signed headers; category and status=available query", "status, timestamp, data[] with eligible IDs", "3 seconds; one retry (two attempts total)", "No remote restriction when URL is not configured"],
    ], [1.25, 1.8, 1.75, 2.2, 1.55, 1.9], font_size=7.2)
    add_paragraph(doc, "ModuleApiClient sorts query parameters before signing, generates a UUID request ID, includes the current ISO-8601 timestamp, and signs the exact request target. It rejects consumed envelopes missing status or timestamp and rejects invalid response timestamps. Laravel's HTTP client directly supports timeouts and retries (Laravel, 2026b).")
    add_code(doc, """private function client(): PendingRequest
{
    return Http::acceptJson()->timeout(3)->retry(2, 100);
}

private function validatedEnvelope(Response $response): array
{
    $payload = $response->json();
    if (! is_array($payload)
        || ! isset($payload['status'], $payload['timestamp'])) {
        throw new UnexpectedValueException(
            'The consumed service returned an invalid IFA envelope.'
        );
    }
    CarbonImmutable::parse((string) $payload['timestamp']);
    return $payload;
}""", "Code excerpt 9. IFA-aware consuming client — app/Services/ModuleApiClient.php")

    add_heading(doc, "6.6 Web-service security and interoperability summary", 2)
    add_table(doc, ["IFA requirement", "Implementation evidence"], [
        ["Request timestamp or request ID mandatory", "Both X-Request-ID and X-Timestamp are mandatory and signed"],
        ["Every response has status and timestamp", "Controller, middleware, validation, and exception responses use the same envelope"],
        ["Expose a service", "POST/GET/DELETE /api/v1/requests endpoints"],
        ["Consume services", "ModuleApiClient calls Module 1 identity and Module 2 donation availability"],
        ["Timeout and retry", "timeout(3)->retry(2, 100): initial call plus one retry"],
        ["Security", "HMAC, clock-skew check, idempotency fingerprint, rate limit, validation, and safe errors"],
    ], [2.4, 5.5], font_size=8.3)

    add_heading(doc, "7. References", 1)
    references = [
        "Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design patterns: Elements of reusable object-oriented software. Addison-Wesley.",
        "Laravel. (2026a). Database: Query builder—Pessimistic locking (Version 13.x). https://laravel.com/framework/docs/13.x/queries#pessimistic-locking",
        "Laravel. (2026b). HTTP client (Version 13.x). https://laravel.com/framework/docs/13.x/http-client",
        "OpenAI. (2026). ChatGPT [Large language model]. https://chatgpt.com/",
        "OWASP Foundation. (2025). OWASP Top 10:2025. https://owasp.org/Top10/",
        "United Nations, Department of Economic and Social Affairs. (n.d.). Goal 2: End hunger, achieve food security and improved nutrition and promote sustainable agriculture. https://sdgs.un.org/goals/goal2",
    ]
    for reference in references:
        p = doc.add_paragraph(reference)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.space_after = Pt(8)

    add_heading(doc, "Appendix A — Requirement Compliance Matrix", 1)
    add_table(doc, ["Assignment requirement", "Evidence in report", "Implementation evidence", "Status"], [
        ["1. Introduction", "Section 1", "System, SDG 2 contribution, target users, scope", "Complete"],
        ["2. Complete module functions and screenshots with class paths", "Section 2; Figures 1–5", "Controllers, validation, matching, views, observers", "Complete"],
        ["3. Entity class diagram; object references, not ERD", "Section 3; Figure 6", "Five Eloquent entity classes with navigable relationships", "Complete"],
        ["4. Unique design pattern; description, class diagram, justification", "Section 4; Figure 7", "Observer interface, publisher, two observers, three outcome events", "Complete"],
        ["5. Two threats and unique secure-coding practices; validation mandatory but not counted", "Section 5", "Transactional row locking; signed/fresh/idempotent module requests; separate validation", "Complete"],
        ["6. Expose and consume web services; request ID/timestamp and response status/timestamp", "Section 6", "Three exposed endpoints; two consumed services; common IFA envelope", "Complete"],
        ["7. APA 7 references and AI disclosure", "Declaration; Section 7", "APA-style list and explicit generative-AI usage disclosure", "Complete"],
    ], [2.2, 1.4, 3.2, 0.8], font_size=7.7)

    add_heading(doc, "Appendix B — Verification Evidence", 1)
    add_heading(doc, "B.1 Automated verification", 2)
    add_table(doc, ["Check", "Result", "Coverage"], [
        ["php artisan test", "PASS — 14 tests, 52 assertions", "Auto matching, API signing, replay handling, client IFA validation, Observer publisher, browser routes"],
        ["vendor/bin/pint --test", "PASS", "Laravel/PHP coding style"],
        ["composer validate --strict", "PASS", "Composer schema and package metadata"],
        ["PHP syntax checks", "PASS", "Changed PHP source files"],
    ], [2.0, 1.9, 4.1], font_size=8.3)
    add_heading(doc, "B.2 End-to-end scenario", 2)
    add_table(doc, ["Step", "Observed result"], [
        ["1. Recipient signs in", "Authenticated recipient dashboard is displayed"],
        ["2. Initial stock", "Packed rice: 40 portions; vegetables: 25 portions"],
        ["3. Recipient requests 12 Packaged Goods", "Request passes validation and automatic matching runs"],
        ["4. Match outcome", "12 of 12 portions allocated; request status matched"],
        ["5. Observer result", "Notification: ‘Match succeeded: 12 of 12 portions allocated.’"],
        ["6. Stock integrity", "Packed rice decreases from 40 to 28"],
        ["7. History", "Matched request and 12-portion allocation are visible"],
        ["8. Signed API call", "HTTP 201 envelope contains requestID, timestamp, status, and data; fingerprint is absent"],
    ], [2.1, 5.9], font_size=8.4)
    add_callout(doc, "Final result", "The verified implementation satisfies the supplied Module 3 scope and the BMIT3173 report requirements for design pattern, entity classes, security, web-service exposure/consumption, IFA fields, screenshots, references, and AI disclosure.", color=GREEN_DARK, fill=GREEN_LIGHT)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
